import customtkinter as ctk
import requests
import threading
import json
from tkinter import messagebox, filedialog
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re

# --- CẤU HÌNH ---
API_URL = "https://api-dan-cu.onrender.com/api"
THEME_COLOR = "#B91C1C" # Đỏ VNeID

ctk.set_appearance_mode("Light")
ctk.set_default_color_theme("blue")

class AppQuanLyDanCu(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("HỆ THỐNG QUẢN LÝ DÂN CƯ (PC)")
        self.geometry("1280x800")
        
        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(0, weight=1)

        # === SIDEBAR (CỘT TRÁI) ===
        self.sidebar = ctk.CTkFrame(self, width=220, corner_radius=0, fg_color=THEME_COLOR)
        self.sidebar.grid(row=0, column=0, sticky="nsew")
        
        # Logo
        ctk.CTkLabel(self.sidebar, text="QUẢN LÝ\nDÂN CƯ", font=ctk.CTkFont(size=24, weight="bold"), text_color="yellow").pack(pady=30)
        
        # Menu
        self.btn_ds = ctk.CTkButton(self.sidebar, text="Danh sách Hộ Dân", command=self.show_danh_sach, fg_color="#7F1D1D", text_color="white", height=45, anchor="w", font=ctk.CTkFont(size=15, weight="bold"))
        self.btn_ds.pack(fill="x", padx=10, pady=5)
        
        self.btn_tk = ctk.CTkButton(self.sidebar, text="Thống kê Số liệu", command=self.show_thong_ke, fg_color="transparent", text_color="white", height=45, anchor="w", font=ctk.CTkFont(size=15, weight="bold"))
        self.btn_tk.pack(fill="x", padx=10, pady=5)

        # Footer Version
        ctk.CTkLabel(self.sidebar, text="Phiên bản 2.1\nDesign by You", text_color="#FECaca", font=ctk.CTkFont(size=12)).pack(side="bottom", pady=20)

        # === MAIN CONTENT (BÊN PHẢI) ===
        self.main_frame = ctk.CTkFrame(self, corner_radius=0, fg_color="#F3F4F6")
        self.main_frame.grid(row=0, column=1, sticky="nsew")
        
        self.frame_danh_sach = ctk.CTkFrame(self.main_frame, fg_color="transparent")
        self.frame_thong_ke = ctk.CTkFrame(self.main_frame, fg_color="transparent")
        
        # Khởi tạo giao diện con
        self.build_ui_danh_sach()
        self.build_ui_thong_ke()
        
        # Mặc định hiện danh sách
        self.show_danh_sach()
        self.data_source = []
        self.load_data()

    # --- CHUYỂN TAB ---
    def show_danh_sach(self):
        self.frame_thong_ke.pack_forget()
        self.frame_danh_sach.pack(fill="both", expand=True, padx=20, pady=20)
        self.btn_ds.configure(fg_color="#7F1D1D")
        self.btn_tk.configure(fg_color="transparent")

    def show_thong_ke(self):
        self.frame_danh_sach.pack_forget()
        self.frame_thong_ke.pack(fill="both", expand=True, padx=20, pady=20)
        self.btn_ds.configure(fg_color="transparent")
        self.btn_tk.configure(fg_color="#7F1D1D")
        self.update_thong_ke()

    # --- UI DANH SÁCH ---
    def build_ui_danh_sach(self):
        # Header
        top = ctk.CTkFrame(self.frame_danh_sach, fg_color="transparent")
        top.pack(fill="x", pady=(0, 15))
        ctk.CTkLabel(top, text="DANH SÁCH HỒ SƠ DÂN CƯ", font=ctk.CTkFont(size=26, weight="bold"), text_color="#1F2937").pack(side="left")
        ctk.CTkButton(top, text="🔄 Tải lại dữ liệu", command=self.load_data, fg_color="#059669", height=40).pack(side="right")

        # Scroll List
        self.scroll_list = ctk.CTkScrollableFrame(self.frame_danh_sach, fg_color="white", corner_radius=10)
        self.scroll_list.pack(fill="both", expand=True)

    def render_list_items(self):
        for w in self.scroll_list.winfo_children(): w.destroy()
        
        if not self.data_source:
            ctk.CTkLabel(self.scroll_list, text="Đang tải hoặc chưa có dữ liệu...", text_color="gray", font=ctk.CTkFont(size=16)).pack(pady=50)
            return

        for idx, item in enumerate(self.data_source):
            try: mems = json.loads(item.get('danh_sach_thanh_vien', '[]')); count = len(mems) + 1
            except: count = 1
            
            # Card Item
            card = ctk.CTkFrame(self.scroll_list, fg_color="#EFF6FF", border_color="#BFDBFE", border_width=1, corner_radius=8)
            card.pack(fill="x", pady=8, padx=10)
            
            # Info Section
            info = ctk.CTkFrame(card, fg_color="transparent")
            info.pack(side="left", padx=15, pady=15, fill="x", expand=True)
            
            name = str(item.get('ho_ten','')).upper()
            ctk.CTkLabel(info, text=f"{idx+1}. {name}", font=ctk.CTkFont(size=18, weight="bold"), text_color="#1E40AF", anchor="w").pack(fill="x")
            
            detail_text = f"🏠 {item.get('thuong_tru','')}  |  👨‍👩‍👧‍👦 {count} thành viên"
            if item.get('nguoi_tao_sdt'): detail_text += f"  |  ✍️ Người nhập: {item.get('nguoi_tao_sdt')}"
            
            ctk.CTkLabel(info, text=detail_text, text_color="#4B5563", font=ctk.CTkFont(size=14), anchor="w").pack(fill="x", pady=(5,0))
            
            # Button Section (CHỈ CÓ NÚT XEM CHI TIẾT)
            ctk.CTkButton(card, text="👁️ Xem Chi Tiết", fg_color="#B91C1C", hover_color="#991B1B", width=140, height=40, font=ctk.CTkFont(weight="bold"),
                          command=lambda d=item: self.open_detail_window(d)).pack(side="right", padx=15)

    # --- CỬA SỔ CHI TIẾT (POPUP) ---
    def open_detail_window(self, data):
        # Tạo cửa sổ mới
        win = ctk.CTkToplevel(self)
        win.title(f"HỒ SƠ: {data.get('ho_ten','').upper()}")
        win.geometry("900x750")
        win.transient(self) # Cửa sổ con của main
        win.grab_set() # Chặn thao tác cửa sổ chính

        # Header Title
        ctk.CTkLabel(win, text="THÔNG TIN CHI TIẾT HỘ DÂN", font=ctk.CTkFont(size=22, weight="bold"), text_color="#B91C1C").pack(pady=(20,10))

        # Scroll Area
        scroll = ctk.CTkScrollableFrame(win, fg_color="white", corner_radius=10)
        scroll.pack(fill="both", expand=True, padx=20, pady=(0, 20))

        # --- PHẦN 1: CHỦ HỘ ---
        ctk.CTkLabel(scroll, text="I. THÔNG TIN CHỦ HỘ", font=ctk.CTkFont(size=18, weight="bold"), text_color="#374151", anchor="w").pack(fill="x", pady=(10,5))
        
        info_grid = ctk.CTkFrame(scroll, fg_color="#F9FAFB", border_color="#E5E7EB", border_width=1)
        info_grid.pack(fill="x", pady=5)

        # Hàm vẽ dòng (Label: Value)
        def add_row(parent, label, value, color="black"):
            row = ctk.CTkFrame(parent, fg_color="transparent")
            row.pack(fill="x", pady=2, padx=10)
            ctk.CTkLabel(row, text=label, width=180, anchor="w", font=ctk.CTkFont(weight="bold", size=14), text_color="#4B5563").pack(side="left")
            ctk.CTkLabel(row, text=str(value), anchor="w", font=ctk.CTkFont(size=14), text_color=color).pack(side="left", fill="x", expand=True)
            ctk.CTkFrame(parent, height=1, fg_color="#E5E7EB").pack(fill="x", padx=10) # Đường kẻ mờ

        add_row(info_grid, "Họ và tên:", data.get('ho_ten', '').upper(), "#B91C1C")
        add_row(info_grid, "Ngày sinh:", data.get('ngay_sinh', ''))
        add_row(info_grid, "Giới tính:", data.get('gioi_tinh', ''))
        add_row(info_grid, "Số CMND/CCCD:", data.get('so_cmnd', ''))
        add_row(info_grid, "Ngày cấp:", data.get('ngay_cap', ''))
        add_row(info_grid, "Nơi cấp:", data.get('noi_cap', ''))
        add_row(info_grid, "Thường trú:", data.get('thuong_tru', ''))
        add_row(info_grid, "Nơi ở hiện tại:", data.get('noi_o_hien_tai', ''))
        add_row(info_grid, "Quê quán:", data.get('que_quan', ''))
        add_row(info_grid, "Dân tộc:", data.get('dan_toc', ''))
        add_row(info_grid, "Tôn giáo:", data.get('ton_giao', ''))
        add_row(info_grid, "Trình độ văn hóa:", data.get('trinh_do', ''))
        add_row(info_grid, "Số điện thoại:", data.get('sdt', ''))
        add_row(info_grid, "Công việc:", data.get('cong_viec', ''))

        # --- PHẦN 2: THÀNH VIÊN ---
        try:
            mems = json.loads(data.get('danh_sach_thanh_vien', '[]'))
            if mems:
                ctk.CTkLabel(scroll, text=f"II. THÀNH VIÊN GIA ĐÌNH ({len(mems)} người)", font=ctk.CTkFont(size=18, weight="bold"), text_color="#374151", anchor="w").pack(fill="x", pady=(20,5))
                
                for idx, m in enumerate(mems):
                    mem_card = ctk.CTkFrame(scroll, fg_color="#EFF6FF", border_color="#BFDBFE", border_width=1)
                    mem_card.pack(fill="x", pady=5)
                    
                    # Header thành viên
                    head = ctk.CTkFrame(mem_card, fg_color="#DBEAFE", height=30)
                    head.pack(fill="x")
                    ctk.CTkLabel(head, text=f"  #{idx+1} - {m.get('ho_ten','').upper()}", font=ctk.CTkFont(weight="bold"), text_color="#1E40AF").pack(side="left")
                    ctk.CTkLabel(head, text=f"Quan hệ: {m.get('quan_he','')}  ", font=ctk.CTkFont(weight="bold"), text_color="#1E40AF").pack(side="right")

                    # Nội dung thành viên
                    mem_content = ctk.CTkFrame(mem_card, fg_color="transparent")
                    mem_content.pack(fill="x", padx=10, pady=5)
                    
                    # Hiển thị vắn tắt thông tin thành viên
                    info_str = f"Sinh: {m.get('ngay_sinh','')} | CMND: {m.get('so_cmnd','')} | Ngày cấp: {m.get('ngay_cap','')} | Nơi cấp: {m.get('noi_cap','')}\n"
                    info_str += f"Dân tộc: {m.get('dan_toc','')} | Tôn giáo: {m.get('ton_giao','')} | Trình độ: {m.get('trinh_do','')}\n"
                    info_str += f"Công việc/Tình trạng: {m.get('cong_viec','')}"
                    
                    # Hiển thị tình trạng checkboxes
                    tt = m.get('tinh_trang', [])
                    if tt: info_str += " | " + (", ".join(tt) if isinstance(tt, list) else str(tt))

                    ctk.CTkLabel(mem_content, text=info_str, justify="left", anchor="w", text_color="#4B5563").pack(fill="x")
            else:
                ctk.CTkLabel(scroll, text="II. THÀNH VIÊN: (Không có)", font=ctk.CTkFont(size=16), text_color="gray", anchor="w").pack(pady=20)

        except Exception as e:
            print("Lỗi hiển thị thành viên:", e)

        # --- FOOTER: NÚT IN ---
        footer = ctk.CTkFrame(win, fg_color="white", height=80)
        footer.pack(fill="x", side="bottom")
        ctk.CTkFrame(footer, height=1, fg_color="#E5E7EB").pack(fill="x") # Đường kẻ
        
        ctk.CTkButton(footer, text="🖨️ XUẤT FILE WORD ĐỂ IN", height=50, width=300, 
                      fg_color="#059669", hover_color="#047857", font=ctk.CTkFont(size=16, weight="bold"),
                      command=lambda: self.export_word(data)).pack(pady=15)

    # --- LOGIC XUẤT WORD (ĐÃ FIX: Ẩn bảng nếu không có thành viên) ---
    def export_word(self, data):
        template_path = "pc_app/mau_phieu.docx"
        if not os.path.exists(template_path): 
            return messagebox.showerror("Lỗi", "Không tìm thấy file 'mau_phieu.docx'")

        try:
            doc = Document(template_path)
            
            # 1. Mapping dữ liệu Chủ hộ
            mapping = {
                "Họ, chữ đệm và tên người khai": data.get('ho_ten', '').upper(),
                "Ngày, tháng, năm sinh": data.get('ngay_sinh', ''),
                "Giới tính": data.get('gioi_tinh', ''),
                "Số CMND": data.get('so_cmnd', ''),
                "Ngày cấp": data.get('ngay_cap', ''),
                "Nơi cấp": data.get('noi_cap', ''),
                "Địa chỉ thường trú": data.get('thuong_tru', ''),
                "Nơi ở hiện tại": data.get('noi_o_hien_tai', ''),
                "Quê quán": data.get('que_quan', ''),
                "Trình độ văn hoá": data.get('trinh_do', ''),
                "Dân tộc": data.get('dan_toc', ''),
                "Tôn giáo": data.get('ton_giao', ''),
                "SĐT": data.get('sdt', ''),
                "Công việc": data.get('cong_viec', '')
            }

            # Hàm replace an toàn (Chỉ thay thế vị trí cần thiết)
            def safe_replace(para):
                text = para.text
                for key, val in mapping.items():
                    if key in text:
                        # Regex: Tìm key + (dấu :) + khoảng trắng + (dấu chấm hoặc 3 chấm)
                        pattern = re.escape(key) + r"(?::)?\s*[.…]{2,}"
                        if val:
                            new_text = re.sub(pattern, f"{key}: {val}", text)
                            if new_text != text: text = new_text
                            else: 
                                if f"{key}: {val}" not in text: text = text.replace(key, f"{key}: {val}")
                para.text = text

            for p in doc.paragraphs: safe_replace(p)

            # 2. Xử lý Thành viên (Tự động thêm nếu có)
            try:
                mems = json.loads(data.get('danh_sach_thanh_vien', '[]'))
                if len(mems) > 0:
                    doc.add_page_break()
                    h = doc.add_paragraph("II. THÔNG TIN NGƯỜI CHUNG HỘ GIA ĐÌNH")
                    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    h.runs[0].bold = True
                    
                    table = doc.add_table(rows=1, cols=6)
                    table.style = 'Table Grid'
                    headers = ["STT", "Họ Tên", "Quan Hệ", "Ngày Sinh", "CMND", "Tình Trạng"]
                    for i, t in enumerate(headers): 
                        run = table.rows[0].cells[i].paragraphs[0].add_run(t)
                        run.bold = True
                    
                    for i, m in enumerate(mems):
                        cells = table.add_row().cells
                        cells[0].text = str(i+1)
                        cells[1].text = m.get('ho_ten', '').upper()
                        cells[2].text = m.get('quan_he', '')
                        cells[3].text = m.get('ngay_sinh', '')
                        cells[4].text = m.get('so_cmnd', '')
                        tt = m.get('tinh_trang', [])
                        cells[5].text = ", ".join(tt) if isinstance(tt, list) else str(tt)
            except: pass

            filename = f"Phieu_{data.get('ho_ten','Noname')}.docx"
            path = filedialog.asksaveasfilename(defaultextension=".docx", initialfile=filename)
            if path:
                doc.save(path)
                os.startfile(path)
                messagebox.showinfo("Thành công", "Đã xuất file Word!")

        except Exception as e: messagebox.showerror("Lỗi", str(e))

    # --- UI THỐNG KÊ ---
    def build_ui_thong_ke(self):
        ctk.CTkLabel(self.frame_thong_ke, text="BẢNG ĐIỀU KHIỂN SỐ LIỆU", font=ctk.CTkFont(size=24, weight="bold"), text_color="#111827").pack(anchor="w", pady=20)
        self.stats_container = ctk.CTkFrame(self.frame_thong_ke, fg_color="transparent")
        self.stats_container.pack(fill="both", expand=True)

    def update_thong_ke(self):
        for w in self.stats_container.winfo_children(): w.destroy()
        if not self.data_source: return
        
        total_ho = len(self.data_source)
        total_nguoi = 0
        addr_map = {}

        for item in self.data_source:
            try: total_nguoi += 1 + len(json.loads(item.get('danh_sach_thanh_vien','[]')))
            except: total_nguoi += 1
            
            addr = item.get('thuong_tru', 'Chưa rõ')
            addr_map[addr] = addr_map.get(addr, 0) + 1

        # Cards
        row1 = ctk.CTkFrame(self.stats_container, fg_color="transparent")
        row1.pack(fill="x", pady=10)
        
        def make_card(p, title, val, color):
            f = ctk.CTkFrame(p, fg_color=color, height=120, corner_radius=10)
            f.pack(side="left", fill="x", expand=True, padx=10)
            ctk.CTkLabel(f, text=title, text_color="white", font=ctk.CTkFont(size=14)).pack(pady=(20,5))
            ctk.CTkLabel(f, text=str(val), text_color="white", font=ctk.CTkFont(size=36, weight="bold")).pack(pady=5)

        make_card(row1, "TỔNG SỐ HỘ", total_ho, "#059669")
        make_card(row1, "TỔNG NHÂN KHẨU", total_nguoi, "#2563EB")

        # Table Detail
        ctk.CTkLabel(self.stats_container, text="CHI TIẾT THEO ĐỊA BÀN", font=ctk.CTkFont(size=16, weight="bold"), text_color="#374151").pack(anchor="w", pady=(30,10))
        
        scroll = ctk.CTkScrollableFrame(self.stats_container, fg_color="white", height=300)
        scroll.pack(fill="x")
        
        for addr, count in addr_map.items():
            r = ctk.CTkFrame(scroll, fg_color="transparent")
            r.pack(fill="x", pady=5)
            ctk.CTkLabel(r, text=addr, font=ctk.CTkFont(size=14)).pack(side="left", padx=10)
            ctk.CTkLabel(r, text=f"{count} hộ", font=ctk.CTkFont(weight="bold")).pack(side="right", padx=10)
            ctk.CTkFrame(scroll, height=1, fg_color="#F3F4F6").pack(fill="x")

    def load_data(self):
        def task():
            try:
                res = requests.get(f"{API_URL}/danh-sach", timeout=60)
                if res.status_code==200:
                    self.data_source = res.json()
                    self.after(0, self.render_list_items)
                else: self.after(0, lambda: messagebox.showerror("Lỗi", "Lỗi tải dữ liệu"))
            except: self.after(0, lambda: messagebox.showerror("Lỗi", "Không kết nối được Server"))
        threading.Thread(target=task).start()

if __name__ == "__main__":
    app = AppQuanLyDanCu()
    app.mainloop()